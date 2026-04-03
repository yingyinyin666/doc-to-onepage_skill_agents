#!/usr/bin/env python3
import argparse
import os
import re
import PyPDF2
import docx

def process_pdf(file_path):
    """处理PDF文件，提取文本内容"""
    text = ""
    try:
        with open(file_path, 'rb') as f:
            reader = PyPDF2.PdfReader(f)
            for page_num in range(len(reader.pages)):
                page = reader.pages[page_num]
                text += page.extract_text() + "\n"
    except Exception as e:
        print(f"Error processing PDF: {e}")
    return text

def process_docx(file_path):
    """处理Word文档，提取文本内容"""
    text = ""
    try:
        doc = docx.Document(file_path)
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + "\t"
                text += "\n"
    except Exception as e:
        print(f"Error processing DOCX: {e}")
    return text

def process_text(file_path):
    """处理文本文件，直接读取内容"""
    try:
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read()
    except Exception as e:
        print(f"Error processing text file: {e}")
        return ""

def process_document(file_path):
    """根据文件类型处理文档，返回提取的文本内容"""
    file_extension = os.path.splitext(file_path)[1].lower()
    
    if file_extension == '.pdf':
        return process_pdf(file_path)
    elif file_extension == '.docx':
        return process_docx(file_path)
    elif file_extension in ['.txt', '.md', '.lark.md']:
        return process_text(file_path)
    else:
        print(f"Unsupported file type: {file_extension}")
        return ""

def analyze_content(text):
    """分析文本内容，提取核心信息"""
    # 提取标题（假设第一行是标题）
    lines = text.strip().split('\n')
    title = lines[0] if lines else ""
    
    # 提取摘要（前几句话）
    summary = " ".join(lines[:3]) if len(lines) >= 3 else " ".join(lines)
    
    # 提取核心要点（假设以列表形式存在）
    key_points = []
    for line in lines:
        if re.match(r'^[\*\-\d\.]+\s+', line):
            key_points.append(line)
    
    # 提取数值数据
    data = []
    for line in lines:
        numbers = re.findall(r'\d+(?:\.\d+)?', line)
        if numbers:
            data.append({"line": line, "numbers": numbers})
    
    return {
        "title": title,
        "summary": summary,
        "key_points": key_points,
        "data": data
    }

def determine_style(content):
    """根据内容性质决定UI风格"""
    text = content.get("summary", "") + " " + " ".join(content.get("key_points", []))
    
    # 关键词匹配
    report_keywords = ["报告", "分析", "数据", "统计", "总结", "汇报"]
    tutorial_keywords = ["教程", "步骤", "指南", "如何", "教程", "代码"]
    promotion_keywords = ["宣传", "推广", "活动", "优惠", "产品", "服务"]
    
    report_count = sum(1 for keyword in report_keywords if keyword in text)
    tutorial_count = sum(1 for keyword in tutorial_keywords if keyword in text)
    promotion_count = sum(1 for keyword in promotion_keywords if keyword in text)
    
    if report_count >= 2:
        return "report"  # 商务蓝、极简风
    elif tutorial_count >= 2:
        return "tutorial"  # 步骤条、代码块
    elif promotion_count >= 2:
        return "promotion"  # 大字报、高饱和度色块
    else:
        return "default"  # 默认风格

def main(file_path, output):
    """主函数"""
    # 处理文档
    text = process_document(file_path)
    
    # 分析内容
    content = analyze_content(text)
    
    # 确定风格
    style = determine_style(content)
    
    # 输出结果
    result = {
        "text": text,
        "content": content,
        "style": style
    }
    
    # 确保输出目录存在
    os.makedirs(os.path.dirname(output), exist_ok=True)
    
    # 写入JSON文件
    import json
    with open(output, 'w', encoding='utf-8') as f:
        json.dump(result, f, ensure_ascii=False, indent=2)
    
    print(f"Document processed successfully: {output}")
    print(f"Determined style: {style}")

if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="Process document and extract content")
    parser.add_argument('--file', type=str, required=True, help='Path to document file')
    parser.add_argument('--output', type=str, default="./output/document_analysis.json", help='Output JSON file path')
    
    args = parser.parse_args()
    main(
        file_path=args.file,
        output=args.output
    )
